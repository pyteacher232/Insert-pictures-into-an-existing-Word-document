# docx module is used for reading and writing a word document.
import docx
# listdir of os is used for getting all file names under a directory.
from os import listdir
# isfile is used for checking if it is a file or directory and join is used for joining path segments.
from os.path import isfile, join

# Path of your word file.
word_fname = "Word Report- Python Automation.docx"
# Path of your pictures.
pic_dir = "pics"

doc = docx.Document(word_fname)

# Get the width of your word document.
width = doc.sections[0].page_width
# Get the left margin of your word document.
left_margin = doc.sections[0].left_margin
# Get the right margin of your word document.
right_margin = doc.sections[0].right_margin

# Calculate the width of images to be inserted pairwise.
img_width = (width - left_margin - right_margin) / 2 * 0.95

# Get all paragraphs.
all_paras = doc.paragraphs

# Get names of the remaining pictures except for pictures like 'Picture1.png', 'Picture2.png'....
remaining_files = [f for f in listdir(pic_dir) if isfile(join(pic_dir, f))]
for i in range(6):
    pic_indexed = "Picture{}.png".format(i + 1)
    # remove Picture1 ~ Picture6 from remaining_files
    if pic_indexed in remaining_files:
        remaining_files.remove(pic_indexed)

for para in all_paras:
    print(para.text)

    # Insert from Picture1 to Picture6.
    for i in range(6):
        if "Insert Picture{} here".format(i + 1) in para.text:
            # Delete "Insert Picture here...".
            para.text = ''

            # Add a picture.
            r = para.add_run()
            r.add_picture('pics/Picture{}.png'.format(i + 1))

    # Insert the remaining pictures except from picture1 to picture6.
    if "Insert 16 pictures pairwise" in para.text:
        # Delete "Insert 16 pictures pairwise".
        para.text = ''
        r = para.add_run()

        # Add all the remaining pictures.
        for p in remaining_files:
            r.add_picture(join(pic_dir, p), width=img_width)

# Save the document
doc.save('Word Report- Python Automation_re.docx')
